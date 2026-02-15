#!/usr/bin/env python3
"""
Convert SOFTWARE_REQUIREMENTS_SPECIFICATION.md to DOCX format
Preserves formatting, headings, tables, and lists
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def parse_markdown_line(line):
    """Parse markdown line and return type and content"""
    # Headers
    if line.startswith('# '):
        return ('heading1', line[2:].strip())
    elif line.startswith('## '):
        return ('heading2', line[3:].strip())
    elif line.startswith('### '):
        return ('heading3', line[4:].strip())
    elif line.startswith('#### '):
        return ('heading4', line[5:].strip())
    elif line.startswith('##### '):
        return ('heading5', line[6:].strip())
    
    # Horizontal rule
    elif line.strip() in ['---', '***', '___']:
        return ('hr', '')
    
    # Unordered list
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        content = line.strip()[2:].strip()
        indent = len(line) - len(line.lstrip())
        level = indent // 2
        return ('bullet', content, level)
    
    # Ordered list
    elif re.match(r'^\d+\.\s+', line.strip()):
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            content = match.group(2)
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            return ('numbered', content, level)
    
    # Bold text check
    elif '**' in line:
        return ('bold_text', line)
    
    # Table row
    elif '|' in line and line.strip().startswith('|'):
        return ('table_row', line)
    
    # Empty line
    elif not line.strip():
        return ('empty', '')
    
    # Code block marker
    elif line.strip().startswith('```'):
        return ('code_block', line.strip())
    
    # Placeholder markers
    elif line.strip().startswith('[Insert ') and line.strip().endswith(']'):
        return ('placeholder', line.strip())
    
    # Normal paragraph
    else:
        return ('paragraph', line.strip())

def add_table_from_markdown(doc, lines, start_idx):
    """Create a table from markdown table lines"""
    # Find all consecutive table rows
    table_lines = []
    idx = start_idx
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if line and not re.match(r'^\|[\s\-:]+\|', line):  # Skip separator rows
            table_lines.append(line)
        idx += 1
    
    if not table_lines:
        return start_idx
    
    # Parse table cells
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells and any(cell for cell in cells):  # Has content
            rows.append(cells)
    
    if not rows:
        return idx
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    return idx

def process_bold_text(paragraph, text):
    """Process text with bold markers and add to paragraph"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # Normal text
            paragraph.add_run(part)

def convert_srs_to_docx(md_file_path, output_path):
    """Convert markdown SRS to DOCX"""
    print(f"Reading {md_file_path}...")
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create document
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    print(f"Converting {len(lines)} lines to DOCX...")
    
    in_code_block = False
    code_lines = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        line_type = parse_markdown_line(line)
        
        # Handle code blocks
        if line_type[0] == 'code_block':
            if in_code_block:
                # End code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue
        
        # Process based on type
        if line_type[0] == 'heading1':
            p = doc.add_heading(line_type[1], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif line_type[0] == 'heading2':
            doc.add_heading(line_type[1], level=2)
        
        elif line_type[0] == 'heading3':
            doc.add_heading(line_type[1], level=3)
        
        elif line_type[0] == 'heading4':
            doc.add_heading(line_type[1], level=4)
        
        elif line_type[0] == 'heading5':
            p = doc.add_paragraph(line_type[1])
            p.style = 'Heading 5'
        
        elif line_type[0] == 'hr':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
        
        elif line_type[0] == 'bullet':
            p = doc.add_paragraph(line_type[1], style='List Bullet')
            if len(line_type) > 2 and line_type[2] > 0:
                p.paragraph_format.left_indent = Inches(0.5 * line_type[2])
        
        elif line_type[0] == 'numbered':
            p = doc.add_paragraph(line_type[1], style='List Number')
            if len(line_type) > 2 and line_type[2] > 0:
                p.paragraph_format.left_indent = Inches(0.5 * line_type[2])
        
        elif line_type[0] == 'table_row':
            # Process table
            i = add_table_from_markdown(doc, lines, i)
            continue
        
        elif line_type[0] == 'placeholder':
            p = doc.add_paragraph()
            run = p.add_run(line_type[1])
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif line_type[0] == 'bold_text':
            p = doc.add_paragraph()
            process_bold_text(p, line_type[1])
        
        elif line_type[0] == 'paragraph' and line_type[1]:
            p = doc.add_paragraph()
            process_bold_text(p, line_type[1])
        
        elif line_type[0] == 'empty':
            pass  # Skip empty lines between paragraphs
        
        i += 1
    
    # Save document
    print(f"Saving to {output_path}...")
    doc.save(output_path)
    
    # Get file size
    file_size = os.path.getsize(output_path)
    print(f"✓ Document created successfully!")
    print(f"  Size: {file_size:,} bytes ({file_size/1024:.2f} KB)")
    print(f"  Location: {output_path}")

if __name__ == "__main__":
    md_file = "Project Documents/SOFTWARE_REQUIREMENTS_SPECIFICATION.md"
    docx_file = "Project Documents/SOFTWARE_REQUIREMENTS_SPECIFICATION.docx"
    
    if os.path.exists(md_file):
        convert_srs_to_docx(md_file, docx_file)
    else:
        print(f"Error: File not found: {md_file}")
